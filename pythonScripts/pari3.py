import os
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement

def add_watermark(document, watermark_path):
    # Get the first page of the document
    first_page = document.sections[0].header.paragraphs[0]

    # Add the watermark image to the header paragraph
    run = first_page.add_run()
    run.add_picture(watermark_path, width=document.sections[0].page_width, height=document.sections[0].page_height)

    # Set the watermark options
    watermark_element = OxmlElement('w:watermark')
    watermark_element.append(OxmlElement('w:docPart'))
    watermark_text_element = OxmlElement('w:watermarkText')
    watermark_text_element.set('xml:space', 'preserve')
    watermark_text_element.text = 'YOUR WATERMARK TEXT HERE'
    watermark_element.append(watermark_text_element)

    # Insert the watermark element at the beginning of the header paragraph
    header_paragraph = first_page.paragraphs[0]
    header_paragraph._element.insert(0, watermark_element)

# Path to the input Word document
input_file = 'myfile.docx'

# Path to the watermark image
current_dir = os.path.dirname(os.path.abspath(__file__))
watermark_image = os.path.join(current_dir, 'mypic.png')

# Load the Word document
doc = Document(input_file)

# Add the watermark to the document
add_watermark(doc, watermark_image)

# Save the modified document with the watermark
output_file = 'output.docx'
doc.save(output_file)
